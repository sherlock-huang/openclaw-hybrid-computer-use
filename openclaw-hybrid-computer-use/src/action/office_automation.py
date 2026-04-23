"""Office 自动化控制器

封装 openpyxl (Excel) 和 python-docx (Word) 操作，
支持无头模式下的表格读写、图表生成、文档编辑和模板填充。
"""

import logging
import os
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

logger = logging.getLogger(__name__)


class ExcelController:
    """Excel 控制器 - 基于 openpyxl"""

    def __init__(self):
        self.workbook: Optional[Any] = None
        self.filepath: Optional[str] = None
        self.logger = logging.getLogger(f"{__name__}.ExcelController")

    def create(self, filepath: str) -> "ExcelController":
        """创建新的 Excel 工作簿"""
        try:
            from openpyxl import Workbook
            self.workbook = Workbook()
            self.filepath = filepath
            self.logger.info(f"创建 Excel: {filepath}")
            return self
        except ImportError:
            raise RuntimeError("openpyxl 未安装，请运行: pip install openpyxl")

    def open(self, filepath: str) -> "ExcelController":
        """打开现有 Excel 文件"""
        try:
            from openpyxl import load_workbook
            path = Path(filepath)
            if not path.exists():
                raise FileNotFoundError(f"Excel 文件不存在: {filepath}")
            self.workbook = load_workbook(filepath)
            self.filepath = filepath
            self.logger.info(f"打开 Excel: {filepath}")
            return self
        except ImportError:
            raise RuntimeError("openpyxl 未安装，请运行: pip install openpyxl")

    def get_sheet(self, sheet_name: Optional[str] = None):
        """获取工作表"""
        if not self.workbook:
            raise RuntimeError("未打开或创建工作簿")
        if sheet_name:
            if sheet_name not in self.workbook.sheetnames:
                self.workbook.create_sheet(title=sheet_name)
            return self.workbook[sheet_name]
        return self.workbook.active

    def read_cell(self, cell: str, sheet_name: Optional[str] = None) -> Any:
        """读取单元格值"""
        sheet = self.get_sheet(sheet_name)
        value = sheet[cell].value
        self.logger.debug(f"读取 [{sheet.title}!{cell}] = {value}")
        return value

    def write_cell(self, cell: str, value: Any, sheet_name: Optional[str] = None) -> "ExcelController":
        """写入单元格值"""
        sheet = self.get_sheet(sheet_name)
        sheet[cell] = value
        self.logger.debug(f"写入 [{sheet.title}!{cell}] = {value}")
        return self

    def read_range(self, range_str: str, sheet_name: Optional[str] = None) -> List[List[Any]]:
        """读取区域值，返回二维列表"""
        sheet = self.get_sheet(sheet_name)
        cells = sheet[range_str]
        if not isinstance(cells, tuple):
            return [[cells.value]]
        if isinstance(cells[0], tuple):
            return [[cell.value for cell in row] for row in cells]
        else:
            return [[cell.value for cell in cells]]

    def write_range(self, start_cell: str, data: List[List[Any]], sheet_name: Optional[str] = None) -> "ExcelController":
        """从起始单元格写入二维数据"""
        sheet = self.get_sheet(sheet_name)
        from openpyxl.utils import get_column_letter
        col_letter = "".join([c for c in start_cell if c.isalpha()])
        row_num = int("".join([c for c in start_cell if c.isdigit()]))
        start_col = 0
        for i, ch in enumerate(reversed(col_letter)):
            start_col += (ord(ch.upper()) - ord("A") + 1) * (26 ** i)

        for r_idx, row_data in enumerate(data):
            for c_idx, value in enumerate(row_data):
                col = get_column_letter(start_col + c_idx)
                cell_addr = f"{col}{row_num + r_idx}"
                sheet[cell_addr] = value
        self.logger.info(f"写入区域 [{sheet.title}!{start_cell}] 数据 {len(data)} 行 x {len(data[0]) if data else 0} 列")
        return self

    def create_chart(
        self,
        chart_type: str,
        data_range: str,
        title: str,
        position: Optional[str] = None,
        sheet_name: Optional[str] = None,
        categories_range: Optional[str] = None,
    ) -> "ExcelController":
        """创建图表并插入到工作表"""
        try:
            from openpyxl.chart import BarChart, LineChart, PieChart, AreaChart, ScatterChart, Reference
        except ImportError:
            raise RuntimeError("openpyxl 未安装，请运行: pip install openpyxl")

        sheet = self.get_sheet(sheet_name)
        chart_map = {
            "bar": BarChart,
            "line": LineChart,
            "pie": PieChart,
            "area": AreaChart,
            "scatter": ScatterChart,
        }
        chart_cls = chart_map.get(chart_type.lower())
        if not chart_cls:
            raise ValueError(f"不支持的图表类型: {chart_type}，支持: {list(chart_map.keys())}")

        chart = chart_cls()
        chart.title = title
        chart.style = 10
        chart.y_axis.title = "数值"
        chart.x_axis.title = "类别"

        data_ref = Reference(sheet, range_string=f"{sheet.title}!{data_range}")
        chart.add_data(data_ref, titles_from_data=True)

        if categories_range:
            cats_ref = Reference(sheet, range_string=f"{sheet.title}!{categories_range}")
            chart.set_categories(cats_ref)

        pos = position or "E5"
        sheet.add_chart(chart, pos)
        self.logger.info(f"创建 {chart_type} 图表 '{title}' @ [{sheet.title}!{pos}]")
        return self

    def set_column_width(self, column: str, width: float, sheet_name: Optional[str] = None) -> "ExcelController":
        """设置列宽"""
        sheet = self.get_sheet(sheet_name)
        sheet.column_dimensions[column].width = width
        return self

    def set_row_height(self, row: int, height: float, sheet_name: Optional[str] = None) -> "ExcelController":
        """设置行高"""
        sheet = self.get_sheet(sheet_name)
        sheet.row_dimensions[row].height = height
        return self

    def save(self, filepath: Optional[str] = None) -> "ExcelController":
        """保存工作簿"""
        if not self.workbook:
            raise RuntimeError("没有打开的工作簿")
        save_path = filepath or self.filepath
        if not save_path:
            raise ValueError("未指定保存路径")
        Path(save_path).parent.mkdir(parents=True, exist_ok=True)
        self.workbook.save(save_path)
        self.logger.info(f"保存 Excel: {save_path}")
        return self

    def close(self) -> None:
        """关闭工作簿"""
        self.workbook = None
        self.filepath = None


class WordController:
    """Word 控制器 - 基于 python-docx"""

    def __init__(self):
        self.document: Optional[Any] = None
        self.filepath: Optional[str] = None
        self.logger = logging.getLogger(f"{__name__}.WordController")

    def create(self, filepath: str) -> "WordController":
        """创建新的 Word 文档"""
        try:
            from docx import Document
            self.document = Document()
            self.filepath = filepath
            self.logger.info(f"创建 Word: {filepath}")
            return self
        except ImportError:
            raise RuntimeError("python-docx 未安装，请运行: pip install python-docx")

    def open(self, filepath: str) -> "WordController":
        """打开现有 Word 文件"""
        try:
            from docx import Document
            path = Path(filepath)
            if not path.exists():
                raise FileNotFoundError(f"Word 文件不存在: {filepath}")
            self.document = Document(filepath)
            self.filepath = filepath
            self.logger.info(f"打开 Word: {filepath}")
            return self
        except ImportError:
            raise RuntimeError("python-docx 未安装，请运行: pip install python-docx")

    def add_heading(self, text: str, level: int = 1) -> "WordController":
        """添加标题"""
        if not self.document:
            raise RuntimeError("未打开或创建文档")
        self.document.add_heading(text, level=level)
        self.logger.debug(f"添加标题 (level={level}): {text}")
        return self

    def add_paragraph(self, text: str = "", style: Optional[str] = None) -> "WordController":
        """添加段落"""
        if not self.document:
            raise RuntimeError("未打开或创建文档")
        self.document.add_paragraph(text, style=style)
        self.logger.debug(f"添加段落: {text[:50]}...")
        return self

    def add_bullet(self, text: str, level: int = 0) -> "WordController":
        """添加项目符号列表项"""
        if not self.document:
            raise RuntimeError("未打开或创建文档")
        p = self.document.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = level * 360000
        p.add_run(text)
        self.logger.debug(f"添加列表项: {text}")
        return self

    def add_numbered(self, text: str, level: int = 0) -> "WordController":
        """添加编号列表项"""
        if not self.document:
            raise RuntimeError("未打开或创建文档")
        p = self.document.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = level * 360000
        p.add_run(text)
        self.logger.debug(f"添加编号项: {text}")
        return self

    def add_table(self, rows: int, cols: int, data: Optional[List[List[str]]] = None) -> "WordController":
        """添加表格"""
        if not self.document:
            raise RuntimeError("未打开或创建文档")
        table = self.document.add_table(rows=rows, cols=cols)
        table.style = "Light Grid Accent 1"
        if data:
            for r_idx, row_data in enumerate(data):
                if r_idx >= rows:
                    break
                row = table.rows[r_idx]
                for c_idx, cell_text in enumerate(row_data):
                    if c_idx >= cols:
                        break
                    row.cells[c_idx].text = str(cell_text)
        self.logger.info(f"添加表格 {rows}x{cols}")
        return self

    def replace_placeholder(self, placeholder: str, replacement: str) -> int:
        """替换文档中的占位符文本"""
        if not self.document:
            raise RuntimeError("未打开或创建文档")
        count = 0
        for paragraph in self.document.paragraphs:
            if placeholder in paragraph.text:
                full_text = paragraph.text
                if paragraph.runs:
                    paragraph.runs[0].text = full_text.replace(placeholder, replacement)
                    for run in paragraph.runs[1:]:
                        run.text = ""
                    count += full_text.count(placeholder)
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        full_text = cell.text
                        if cell.paragraphs and cell.paragraphs[0].runs:
                            cell.paragraphs[0].runs[0].text = full_text.replace(placeholder, replacement)
                            for run in cell.paragraphs[0].runs[1:]:
                                run.text = ""
                            count += full_text.count(placeholder)
        self.logger.info(f"替换占位符 '{placeholder}' -> '{replacement[:50]}...' (共 {count} 处)")
        return count

    def fill_template(self, mapping: Dict[str, str]) -> int:
        """批量填充模板占位符"""
        count = 0
        for placeholder, replacement in mapping.items():
            count += self.replace_placeholder(placeholder, replacement)
        return count

    def add_page_break(self) -> "WordController":
        """添加分页符"""
        if not self.document:
            raise RuntimeError("未打开或创建文档")
        self.document.add_page_break()
        return self

    def save(self, filepath: Optional[str] = None) -> "WordController":
        """保存文档"""
        if not self.document:
            raise RuntimeError("没有打开的文档")
        save_path = filepath or self.filepath
        if not save_path:
            raise ValueError("未指定保存路径")
        Path(save_path).parent.mkdir(parents=True, exist_ok=True)
        self.document.save(save_path)
        self.logger.info(f"保存 Word: {save_path}")
        return self

    def close(self) -> None:
        """关闭文档"""
        self.document = None
        self.filepath = None


class OfficeAutomationManager:
    """Office 自动化管理器 - 统一管理 Excel/Word 实例"""

    def __init__(self):
        self.excel: Optional[ExcelController] = None
        self.word: Optional[WordController] = None
        self.logger = logging.getLogger(f"{__name__}.OfficeAutomationManager")

    def excel_create(self, filepath: str) -> ExcelController:
        self.excel = ExcelController().create(filepath)
        return self.excel

    def excel_open(self, filepath: str) -> ExcelController:
        self.excel = ExcelController().open(filepath)
        return self.excel

    def word_create(self, filepath: str) -> WordController:
        self.word = WordController().create(filepath)
        return self.word

    def word_open(self, filepath: str) -> WordController:
        self.word = WordController().open(filepath)
        return self.word

    def close_all(self):
        if self.excel:
            self.excel.close()
            self.excel = None
        if self.word:
            self.word.close()
            self.word = None

